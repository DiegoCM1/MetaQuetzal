"""
Genera el reporte técnico de Alerts/SIAT-CT/SMN como archivo .docx
Ejecutar: python3 generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos globales ──────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)

set_margins(doc)

def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8A)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8A)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_parts is None:
        p.add_run(text)
    else:
        # bold_parts: list of (substr, is_bold)
        for chunk, is_bold in bold_parts:
            r = p.add_run(chunk)
            r.bold = is_bold
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after  = Pt(3)
    p.add_run(text)
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # light gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    return p

def separator():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A568A")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1A568A")
        tcPr.append(shd)
    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "DDEEFF" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("REPORTE TÉCNICO")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8A)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run("Módulo Alerts / SIAT-CT / SMN")
run2.bold = True
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
p_proj = doc.add_paragraph()
p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_proj.add_run("BluEye Backend — API v1").bold = True

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date.add_run(f"Fecha: {datetime.date.today().strftime('%d de %B de %Y')}")

p_author = doc.add_paragraph()
p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_author.add_run("Autor: Edgar (Boro739)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. VISIÓN GENERAL
# ═══════════════════════════════════════════════════════════════════════════════
h1("1. Visión General")
separator()

body(
    "El módulo de Alerts es la capa de presentación de riesgo del sistema BluEye. "
    "No genera datos propios: integra, transforma y prioriza información proveniente "
    "de tres fuentes independientes para entregar al usuario móvil una respuesta "
    "cohesiva y personalizada sobre el estado de amenaza ciclónica en su ubicación."
)

h3("Relación con otros módulos")

body("Relación con users:", bold_parts=[
    ("Relación con users: ", True),
    ("El endpoint principal (/alerts/active) es consciente del usuario autenticado. "
     "Obtiene el perfil del usuario desde la feature users — específicamente sus "
     "coordenadas lat/lon — para personalizar el riesgo SIAT. Si el usuario no tiene "
     "perfil registrado, la respuesta sigue siendo válida pero omite el componente personalizado.", False)
])

body("", bold_parts=[
    ("Relación con notifications: ", True),
    ("El módulo alerts no envía notificaciones directamente. Esa responsabilidad "
     "pertenece a siat/service.py, que usa notifications/service.get_tokens_for_users() "
     "para obtener los tokens FCM de los usuarios escalados y enviarles push de forma "
     "dirigida (uno a uno, no broadcast).", False)
])

h3("Proveedores externos")

add_table(
    ["Proveedor", "Qué aporta", "Cómo se accede"],
    [
        ["NHC (National Hurricane Center)", "Posición, intensidad y velocidad de ciclones activos", "JSON público, sin autenticación"],
        ["OpenWeather (OneCall 3.0)", "Condiciones actuales + alertas meteorológicas por coordenada", "REST API con API key"],
        ["SMN/CONAGUA", "Boletín meteorológico nacional diario", "Scraping HTML del sitio oficial"],
    ],
    col_widths=[1.8, 2.8, 2.2],
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. ALERTS
# ═══════════════════════════════════════════════════════════════════════════════
h1("2. Alerts")
separator()

h3("Qué hace el módulo")
body(
    "El módulo alerts tiene dos responsabilidades principales: "
    "mantener un registro de alertas del sistema (creadas manualmente por un administrador) "
    "y actuar como punto de entrada unificado para todo el consumo de alertas por parte del "
    "cliente móvil, a través del endpoint /alerts/active."
)

h3("Endpoints disponibles")

add_table(
    ["Método", "Ruta", "Auth", "Propósito"],
    [
        ["POST",  "/api/v1/alerts",            "X-Api-Key",      "Crear alerta de sistema (uso interno/admin)"],
        ["GET",   "/api/v1/alerts",            "Firebase Bearer", "Listar alertas del sistema (paginado)"],
        ["GET",   "/api/v1/alerts/{id}",       "Firebase Bearer", "Detalle de una alerta específica"],
        ["GET",   "/api/v1/alerts/active",     "Firebase Bearer", "Feed unificado personalizado por usuario"],
        ["GET",   "/api/v1/alerts/weather",    "Firebase Bearer", "Clima actual + alertas por coordenada"],
        ["GET",   "/api/v1/alerts/weather/mx", "Firebase Bearer", "Panorama meteorológico de 12 puntos de México"],
    ],
    col_widths=[0.7, 2.2, 1.5, 2.4],
)

h3("Datos que devuelve")
body(
    "Las alertas del sistema (AlertSummary) exponen: identificador UUID, timestamp de creación, "
    "nivel de severidad 1–5 en escala SIAT, puntuación de riesgo, título y descripción corta. "
    "El detalle (AlertDetail) añade coordenadas geográficas, lista de factores de riesgo y "
    "recomendaciones en formato JSON."
)

h3("Cómo se construyen las alertas de sistema")
body(
    "Las alertas son creadas a través del endpoint POST /api/v1/alerts con autenticación por "
    "API key. El nivel aceptado es entero de 1 a 5, mapeado directamente a la escala SIAT-CT. "
    "Se persisten en PostgreSQL con gen_random_uuid() como identificador."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. /api/v1/alerts/active
# ═══════════════════════════════════════════════════════════════════════════════
h1("3. Endpoint /api/v1/alerts/active")
separator()

h3("Por qué es el endpoint principal")
body(
    "Este es el endpoint que consume la pantalla principal de la app móvil. En lugar de "
    "obligar al cliente a hacer tres llamadas separadas (SIAT, alertas del sistema, boletín SMN), "
    "el backend las agrega en una sola respuesta estructurada con prioridad explícita. "
    "Esto reduce latencia del cliente, centraliza la lógica de priorización y mantiene el frontend simple."
)

h3("Campos de la respuesta (ActiveAlertsResponse)")

body("generated_at:", bold_parts=[
    ("generated_at:  ", True),
    ("Timestamp UTC del momento exacto en que el servidor construyó la respuesta. "
     "Permite al cliente saber qué tan reciente es la información y detectar respuestas cacheadas.", False)
])
body("user_siat:", bold_parts=[
    ("user_siat:  ", True),
    ("Estado SIAT personalizado del usuario autenticado. Contiene nivel actual (1–5), "
     "color correspondiente, razón detallada (nombre del ciclón, distancia, vientos, ETA) "
     "y fecha de última actualización. Es null si el usuario no tiene perfil con coordenadas.", False)
])
body("cyclonic_alerts:", bold_parts=[
    ("cyclonic_alerts:  ", True),
    ("Lista de alertas accionables derivadas del SIAT. Solo se popula cuando el nivel es "
     "VERDE (2) o superior. Nivel AZUL (1) es solo informativo y no genera alerta visible al usuario.", False)
])
body("general_alerts:", bold_parts=[
    ("general_alerts:  ", True),
    ("Últimas 5 alertas creadas por administración del sistema. Misma estructura que las "
     "alertas ciclónicas pero con fuente SYSTEM.", False)
])
body("smn_bulletin:", bold_parts=[
    ("smn_bulletin:  ", True),
    ("Boletín más reciente del SMN. Contiene número de aviso, fecha/hora de emisión, "
     "titular, resumen del pronóstico y enlace al PDF oficial. Puede ser null sin afectar los otros campos.", False)
])

h3("Prioridad de las fuentes")
code_block(
    "SIAT-CT  (riesgo ciclónico personalizado)   ← fuente primaria\n"
    "     ↓\n"
    "Alertas de sistema  (alertas admin)          ← fuente secundaria\n"
    "     ↓\n"
    "Boletín SMN  (información meteorológica)     ← fuente complementaria, best-effort"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. SIAT-CT
# ═══════════════════════════════════════════════════════════════════════════════
h1("4. SIAT-CT")
separator()

h3("Qué es SIAT")
body(
    "El Sistema de Alerta Temprana para Ciclones Tropicales (SIAT-CT) es el sistema oficial "
    "de México para comunicar el nivel de riesgo ciclónico a la población. Define cinco niveles "
    "de alerta, cada uno con un color y un tiempo de anticipación estimado al impacto."
)

add_table(
    ["Nivel", "Color", "Tiempo estimado al impacto"],
    [
        ["1", "AZUL",      "Aviso preventivo — más de 72 horas"],
        ["2", "VERDE",     "Preparación — 24 a 72 horas"],
        ["3", "AMARILLO",  "Alerta — 12 a 24 horas"],
        ["4", "NARANJA",   "Peligro alto — 6 a 12 horas"],
        ["5", "ROJO",      "Impacto inminente — menos de 6 horas"],
    ],
    col_widths=[0.7, 1.2, 4.8],
)

body(
    "No existe una API oficial del SIAT-CT. BluEye deriva estos niveles localmente "
    "a partir de datos reales del NHC (National Hurricane Center)."
)

h3("Capas de la implementación")

body("Proveedor NHC (siat/providers/nhc.py):", bold_parts=[
    ("Proveedor NHC (siat/providers/nhc.py):  ", True),
    ("Consulta el feed público CurrentStorms.json del NHC. "
     "Por cada tormenta normaliza los campos: convierte viento de nudos a km/h (× 1.852), "
     "velocidad de movimiento de mph a km/h (× 1.60934), y extrae posición, presión, "
     "dirección de movimiento y clasificación (TD/TS/HU).", False)
])
body("Evaluador (siat/evaluator.py):", bold_parts=[
    ("Evaluador (siat/evaluator.py):  ", True),
    ("Recibe la posición del usuario y los datos del ciclón, y calcula el nivel de riesgo. "
     "Es una función pura, sin estado ni acceso a base de datos.", False)
])
body("Servicio-orquestador (siat/service.py):", bold_parts=[
    ("Servicio-orquestador (siat/service.py):  ", True),
    ("Coordina el ciclo completo: obtiene ciclones del NHC, obtiene usuarios con ubicación, "
     "ejecuta el evaluador para cada par usuario-ciclón, persiste resultados, determina "
     "escalaciones y envía notificaciones push.", False)
])

h3("Cálculo del riesgo — Estrategia ETA (caso normal)")
body(
    "Primero calcula la distancia entre el usuario y el ciclón con la fórmula de Haversine, "
    "que considera la curvatura de la Tierra. Luego calcula el tiempo estimado de llegada:"
)
code_block("ETA (horas) = distancia_km / velocidad_movimiento_kmh")

body("El nivel SIAT se determina directamente del ETA:")
code_block(
    "ETA < 6h      →  ROJO    (5)\n"
    "ETA 6 - 12h   →  NARANJA (4)\n"
    "ETA 12 - 24h  →  AMARILLO(3)\n"
    "ETA 24 - 72h  →  VERDE   (2)\n"
    "ETA > 72h     →  AZUL    (1)"
)

h3("Casos especiales")

body("Ciclón lejano (más de 1,500 km):", bold_parts=[
    ("Ciclón lejano (más de 1,500 km):  ", True),
    ("El evaluador retorna con out_of_range = True. El servicio detecta este flag "
     "y omite completamente la escritura en DB y el envío de notificaciones. "
     "Más allá de 1,500 km, incluso a 20 km/h, el ciclón tardaría > 75 horas — "
     "los assessments serían ruido operativo sin valor.", False)
])

body("Ciclón estacionario (velocidad = 0):", bold_parts=[
    ("Ciclón estacionario (velocidad = 0):  ", True),
    ("El ETA es matemáticamente indefinido. Se usa una estrategia de respaldo basada "
     "en distancia e intensidad, con umbrales más conservadores para reflejar el "
     "peligro sostenido de un ciclón estancado:", False)
])
code_block(
    "< 150 km  Y  vientos ≥ 150 km/h  →  ROJO    (5)\n"
    "< 300 km  Y  vientos ≥ 100 km/h  →  NARANJA (4)\n"
    "< 500 km  Y  vientos ≥  75 km/h  →  AMARILLO(3)\n"
    "< 700 km                          →  VERDE   (2)\n"
    "Resto                             →  AZUL    (1)"
)

body("Ciclón rápido acercándose:", bold_parts=[
    ("Ciclón rápido acercándose:  ", True),
    ("No requiere lógica especial. Un ciclón a 68 km moviéndose a 30 km/h tiene "
     "ETA = 2.3 horas → nivel ROJO (5). La fórmula ETA maneja esto naturalmente.", False)
])

h3("Variables que usa el evaluador")
add_table(
    ["Variable", "Fuente", "Uso"],
    [
        ["lat/lon del usuario",         "Tabla users (guardada por la app)",              "Punto de referencia para Haversine"],
        ["lat/lon del ciclón",          "NHC latitude_numeric / longitude_numeric",       "Posición actual de la tormenta"],
        ["wind_kmh",                    "NHC intensity (nudos) × 1.852",                  "Nivel en ruta estacionaria"],
        ["movement_speed_kmh",          "NHC movementSpeed (mph) × 1.60934",              "Elige estrategia: ETA vs. distancia"],
        ["distance_km",                 "Calculada con Haversine",                        "Base de ambas estrategias"],
        ["eta_hours",                   "distance_km / movement_speed_kmh",               "Determina nivel en ruta ETA"],
    ],
    col_widths=[1.5, 2.5, 2.7],
)

h3("Ciclo automatizado")
body(
    "El ciclo SIAT se ejecuta como tarea asyncio en background, iniciada en el lifespan de "
    "FastAPI. Corre inmediatamente al arrancar el servidor y luego cada 30 minutos. "
    "Las notificaciones push solo se envían cuando el nivel del usuario sube respecto al "
    "ciclo anterior. El nivel mínimo para notificar es VERDE (2); el nivel AZUL es solo informativo."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. BOLETÍN SMN
# ═══════════════════════════════════════════════════════════════════════════════
h1("5. Boletín SMN / CONAGUA")
separator()

h3("Cómo se obtiene")
body(
    "El Servicio Meteorológico Nacional no dispone de API pública ni RSS feed funcional para "
    "su pronóstico diario. BluEye obtiene esta información mediante scraping HTML del sitio oficial. "
    "El cliente HTTP usa httpx con verificación SSL deshabilitada (verify=False) porque el sitio "
    "del SMN presenta certificados con problemas intermitentes, y con follow_redirects=True para "
    "manejar redirecciones del CMS gubernamental."
)

h3("Datos que se extraen")
add_table(
    ["Campo", "Selector en HTML", "Descripción"],
    [
        ["aviso_num",  "noaviso_valor",                "Número único del boletín del día"],
        ["headline",   "sintesis → <strong>",          "Titular del pronóstico (campo obligatorio)"],
        ["issued_at",  "DD de Mes del AAAA + HH:MMh",  "Fecha y hora de emisión como ISO 8601 UTC"],
        ["summary",    "visforms_contenedor_cuerpo",    "Primer párrafo del pronóstico (máx. 500 chars)"],
        ["pdf_url",    "href con dominio smn.conagua",  "Enlace de descarga directa al PDF oficial"],
    ],
    col_widths=[1.0, 2.2, 3.5],
)

h3("Integración con Alerts")
body(
    "El boletín se obtiene en /alerts/active como la última de las tres fuentes. El campo "
    "smn_bulletin en la respuesta es opcional y puede ser null. El fetch y parse están envueltos "
    "en try/except: cualquier error es capturado, registrado como warning en logs, y la "
    "respuesta continúa normalmente. El usuario no percibe el fallo."
)

h3("Sistema de caché")
body(
    "Para evitar scraping en cada request, el boletín se cachea en memoria del proceso con TTL "
    "de 30 minutos. Si la fetch de red falla, se devuelve el dato cacheado anterior (stale cache), "
    "garantizando que un fallo temporal del SMN no resulte en ausencia de información. "
    "El caché se pierde si el servidor se reinicia."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. WEATHER
# ═══════════════════════════════════════════════════════════════════════════════
h1("6. Weather — /weather y /weather/mx")
separator()

h3("/api/v1/alerts/weather")
body(
    "Recibe coordenadas geográficas (lat, lon) como query parameters y devuelve las condiciones "
    "meteorológicas actuales más las alertas activas para esa ubicación, usando la OneCall 3.0 API "
    "de OpenWeather."
)
body(
    "Antes de llamar a OpenWeather, el servicio valida que las coordenadas estén dentro de los "
    "límites geográficos de México (lat 14–33°N, lon 118.5–86°W). Si están fuera, retorna 422 "
    "con mensaje descriptivo."
)
body(
    "La respuesta incluye: posición, zona horaria, condiciones actuales (temperatura, humedad, "
    "presión, viento), nivel de riesgo Saffir-Simpson calculado desde la velocidad del viento, "
    "y lista de alertas activas de OpenWeather."
)

h3("/api/v1/alerts/weather/mx")
body(
    "Provee un panorama meteorológico nacional para 12 puntos de referencia distribuidos en México: "
    "Tijuana, Monterrey, Chihuahua, Guadalajara, Ciudad de México, Veracruz, Acapulco, Mérida, "
    "Cancún, Villahermosa, Oaxaca y La Paz."
)
body(
    "La implementación realiza las 12 llamadas a OpenWeather en paralelo usando asyncio.gather, "
    "reduciendo el tiempo de respuesta de ~12 segundos secuenciales a ~1–2 segundos concurrentes."
)

h3("Manejo de errores parciales en /weather/mx")
body(
    "Se usa return_exceptions=True en el asyncio.gather. Si una o varias de las 12 llamadas falla, "
    "las excepciones son capturadas como valores en lugar de abortar todo el gather. El servicio "
    "filtra los resultados: los que son instancias de Exception generan un log de warning y se omiten; "
    "los exitosos se incluyen en la respuesta. El cliente recibe los puntos disponibles con un "
    "points_count ajustado, en lugar de recibir un 502 por un fallo parcial."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 7. RESILIENCIA
# ═══════════════════════════════════════════════════════════════════════════════
h1("7. Resiliencia del Sistema")
separator()

add_table(
    ["Escenario de falla", "Comportamiento", "Impacto en usuario"],
    [
        ["SMN no responde (red)",
         "fetch_latest_bulletin retorna caché anterior (stale)",
         "smn_bulletin muestra datos de hasta 30 min de antigüedad"],
        ["SMN cambia su HTML",
         "_parse_bulletin retorna None; WARNING en logs",
         "smn_bulletin = null en /active; resto sin cambio"],
        ["OpenWeather falla en 1 punto (/weather/mx)",
         "return_exceptions captura la excepción; punto omitido",
         "points_count reducido; no hay 502"],
        ["OpenWeather falla en todos los puntos",
         "Lista vacía, points_count = 0",
         "Respuesta 200 con array vacío"],
        ["Sin ciclones activos (NHC)",
         "run_cycle retorna cyclones_found = 0; sin escrituras en DB",
         "cyclonic_alerts = [] en /active"],
        ["Usuario sin perfil (sin lat/lon)",
         "Bloque SIAT omitido completamente",
         "user_siat = null; alertas sistema y SMN sin cambio"],
    ],
    col_widths=[1.8, 2.5, 2.4],
)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. FLUJO END-TO-END
# ═══════════════════════════════════════════════════════════════════════════════
h1("8. Flujo Completo End-to-End")
separator()

h3("Request del usuario → /alerts/active")

steps = [
    ("1. Usuario abre la app",
     "La app envía GET /api/v1/alerts/active con el token Firebase en Authorization: Bearer."),
    ("2. Autenticación",
     "firebase_admin.auth.verify_id_token() valida el token. Si es inválido: 401. Si es válido: extrae uid."),
    ("3. Resolución de perfil",
     "get_user_by_firebase_uid() hace SELECT en la tabla users con el uid de Firebase."),
    ("4. Estado SIAT",
     "get_user_siat_state() ejecuta un JOIN entre user_alert_states y siat_assessments. "
     "Devuelve nivel, color y razón del último ciclo evaluado (pre-calculado en background)."),
    ("5. Alertas del sistema",
     "SELECT en la tabla alerts, últimas 5 filas ordenadas por timestamp DESC."),
    ("6. Boletín SMN",
     "fetch_latest_bulletin() devuelve caché vigente o hace fetch real al SMN."),
    ("7. Construcción de respuesta",
     "El router ensambla ActiveAlertsResponse con generated_at, los tres bloques de datos "
     "y la lógica de nivel >= 2 para cyclonic_alerts."),
]

for title, desc in steps:
    body("", bold_parts=[(title + ": ", True), (desc, False)])

h3("Ciclo SIAT en background (paralelo al request del usuario)")

bg_steps = [
    "El background loop corre asyncio.create_task al iniciar la app.",
    "Cada 30 minutos (y al arrancar): fetch_active_cyclones() consulta NHC CurrentStorms.json.",
    "Por cada ciclón activo: SELECT de todos los usuarios con lat/lon en DB.",
    "evaluate_user() calcula nivel, color, ETA y razón para cada par usuario-ciclón.",
    "Si out_of_range = True: skip. No se escribe en DB, no se envía notificación.",
    "UPSERT en siat_assessments y user_alert_states con el resultado.",
    "Si nivel nuevo > nivel anterior (escalación) Y nivel >= 2: el usuario entra a la cola de push.",
    "get_tokens_for_users() obtiene tokens FCM. messaging.send_each_for_multicast() envía.",
]
for s in bg_steps:
    bullet(s)

# ═══════════════════════════════════════════════════════════════════════════════
# 9. LIMITACIONES
# ═══════════════════════════════════════════════════════════════════════════════
h1("9. Limitaciones Actuales")
separator()

h3("Dirección del ciclón no considerada")
body(
    "El cálculo de ETA usa distancia / velocidad sin considerar hacia dónde se mueve el ciclón. "
    "Un ciclón alejándose produce el mismo ETA que uno acercándose. La implementación resuelve "
    "esto optando siempre por sobre-alertar: es preferible una falsa alarma que un aviso perdido. "
    "Esta limitación está documentada en el header de evaluator.py. Se resolvería incorporando "
    "movementDir del NHC y cálculo de bearing geográfico en una versión futura."
)

h3("Dependencia del HTML del SMN")
body(
    "El boletín SMN depende de patrones regex sobre el HTML de un sitio gubernamental sin API pública. "
    "Si CONAGUA rediseña el sitio, los selectores pueden dejar de funcionar. "
    "El impacto está contenido: la app no muestra boletín hasta actualizar el selector. "
    "Los logs producen warnings descriptivos cuando esto ocurre."
)

h3("Caché SMN en memoria")
body(
    "El caché de 30 minutos vive en memoria del proceso. Un reinicio del servidor (deploy en Railway) "
    "vacía el caché. En ese caso, la primera solicitud hace fetch real al SMN. "
    "Para alta frecuencia de deploys, podría migrarse a Redis."
)

h3("Ciclo SIAT sin scheduler externo")
body(
    "El background loop corre dentro del proceso uvicorn. Con una instancia única funciona "
    "correctamente. Si Railway escala a múltiples instancias, cada una correría su propio ciclo "
    "generando evaluaciones duplicadas. Para este MVP con una sola instancia es el diseño correcto."
)

h3("Sin historial de trayectoria ciclónica")
body(
    "La tabla cyclone_events tiene un constraint UNIQUE (source, name), lo que significa que cada "
    "tormenta tiene una sola fila que se actualiza con su posición más reciente. No existe historial "
    "de trayectoria. Para visualización de ruta se necesitaría un diseño temporal diferente."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 10. CONCLUSIÓN
# ═══════════════════════════════════════════════════════════════════════════════
h1("10. Conclusión")
separator()

h3("Por qué la implementación es válida")
body(
    "La decisión de derivar los niveles SIAT-CT localmente a partir del feed NHC, en lugar de "
    "depender de una API oficial inexistente, es la única solución técnicamente viable. "
    "La fórmula de ETA sobre datos de movimiento real del NHC refleja fielmente la definición "
    "oficial del SIAT-CT. La escala de 5 niveles implementada es la escala oficial y los "
    "umbrales de tiempo son los publicados por CENAPRED."
)

body(
    "La arquitectura de tres fuentes en /alerts/active (SIAT personalizado + alertas del sistema "
    "+ boletín SMN) es exactamente lo que necesita la pantalla principal de una app de emergencias: "
    "riesgo específico para el usuario, información de administración, y contexto meteorológico, "
    "todo en un solo round-trip."
)

h3("Evidencia de funcionamiento en producción")
add_table(
    ["Verificación", "Resultado"],
    [
        ["18 tests automatizados (alerts + siat)",    "18/18 PASS"],
        ["SMN boletín en vivo (25 Abr 2026)",         "Aviso #229 obtenido correctamente"],
        ["NHC feed en vivo",                          "0 tormentas activas (correcto, temporada baja)"],
        ["DB Neon desde servidor local",              "health-db: OK, reachable"],
        ["SIAT run-cycle con API key",                "200 OK, cyclones_found: 0"],
        ["Auth guards (sin token / token inválido)",  "422 / 401 según corresponde"],
        ["Evaluador: ciclón a 7,546 km",              "out_of_range=True, no escribe en DB"],
        ["Evaluador: Cat-4 estacionario a 94 km",     "ROJO (5), correcto"],
    ],
    col_widths=[3.5, 3.2],
)

h3("Valor que aporta al proyecto BluEye")
body(
    "Este módulo transforma tres fuentes de datos dispares — un feed JSON del gobierno de EE.UU., "
    "una API meteorológica comercial, y un sitio HTML del gobierno mexicano — en una respuesta única, "
    "consistente, resiliente y personalizada por usuario. La app móvil no necesita saber cómo "
    "funciona el NHC, ni qué es el SIAT-CT, ni cómo parsear el HTML del SMN. Esa complejidad está "
    "encapsulada en el backend, y el cliente consume una abstracción limpia y contractualmente estable."
)

h3("Veredicto")
p = doc.add_paragraph()
run = p.add_run(
    "La implementación cumple como MVP profesional. Los endpoints responden correctamente, "
    "la lógica de riesgo es coherente con la definición oficial del SIAT-CT, el sistema es "
    "resiliente a fallos de terceros, y está respaldado por 18 tests automatizados que cubren "
    "los casos de uso principales y los edge cases críticos."
)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8A)

# ── Footer ────────────────────────────────────────────────────────────────────
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run("BluEye Backend — Módulo Alerts / SIAT-CT / SMN — Confidencial").font.size = Pt(8)

# ── Guardar ───────────────────────────────────────────────────────────────────
out = "/home/edgar/BlueAi/BluEye/Reporte_Tecnico_Alerts_SIAT_SMN.docx"
doc.save(out)
print(f"Archivo generado: {out}")
